import os
import re
import uuid
import base64
import io
import json
import requests
import fitz
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
import config
import traceback

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ================= 百度 OCR 服务 =================
def get_baidu_ocr_token():
    """获取百度 OCR 的 access_token"""
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {
        "grant_type": "client_credentials",
        "client_id": config.BAIDU_OCR_API_KEY,
        "client_secret": config.BAIDU_OCR_SECRET_KEY,
    }
    try:
        response = requests.post(url, params=params, timeout=10)
        print(f"[DEBUG] 百度Token响应状态: {response.status_code}")
        result = response.json()
        print(f"[DEBUG] 百度Token响应: {result}")
        if "access_token" in result:
            print("[DEBUG] ✅ 获取百度Token成功")
            return result.get("access_token")
        else:
            print(f"[ERROR] 获取百度Token失败: {result}")
            return None
    except Exception as e:
        print(f"[ERROR] 百度Token请求异常: {e}")
        return None

def call_baidu_ocr(image_base64):
    """调用百度通用文字识别"""
    token = get_baidu_ocr_token()
    if not token:
        print("[ERROR] ❌ 无法获取百度Token")
        return {"error": "无法获取百度Token，请检查API Key和Secret Key"}
    
    url = f"https://aip.baidubce.com/rest/2.0/ocr/v1/general_basic?access_token={token}"
    payload = {"image": image_base64}
    headers = {"content-type": "application/x-www-form-urlencoded"}
    
    try:
        response = requests.post(url, data=payload, headers=headers, timeout=30)
        print(f"[DEBUG] 百度OCR响应状态: {response.status_code}")
        result = response.json()
        print(f"[DEBUG] 百度OCR响应内容: {str(result)[:500]}...")  # 只打印前500字符
        return result
    except Exception as e:
        print(f"[ERROR] 百度OCR调用异常: {e}")
        return {"error": str(e)}

# ================= DeepSeek API 服务 =================
def call_deepseek(prompt):
    """调用 DeepSeek API 生成内容"""
    url = f"{config.DEEPSEEK_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {config.DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是一个专业的财务助手，帮助用户生成规范的报销单内容。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }
    try:
        response = requests.post(url, json=data, headers=headers, timeout=60)
        result = response.json()
        return result.get("choices", [{}])[0].get("message", {}).get("content", "")
    except Exception as e:
        print(f"[ERROR] DeepSeek API调用失败: {e}")
        return None

# ================= 前端 API 接口 =================
@app.get("/")
def read_root():
    return {"message": "数智办公 - 报销会计后端服务已启动"}

@app.post("/api/ocr")
async def ocr_invoice(file: UploadFile = File(...)):
    """接收图片/PDF，调用百度 OCR 识别"""
    print(f"[DEBUG] ===== 收到OCR请求 =====")
    print(f"[DEBUG] 文件名: {file.filename}")
    print(f"[DEBUG] 文件类型: {file.content_type}")
    
    try:
        # 读取文件
        image_data = await file.read()
        print(f"[DEBUG] 文件大小: {len(image_data)} 字节")
        
        if len(image_data) == 0:
            print("[ERROR] ❌ 文件为空")
            return JSONResponse(
                status_code=400,
                content={"code": 400, "detail": "上传的文件为空，请重新上传"}
            )
        
        # 处理 PDF 文件：转图片再 OCR
        if file.content_type == 'application/pdf' or file.filename.lower().endswith('.pdf'):
            print("[DEBUG] 检测到PDF文件，转换为图片...")
            try:
                pdf_doc = fitz.open(stream=image_data, filetype="pdf")
                # 取第一页
                page = pdf_doc[0]
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                image_base64 = base64.b64encode(img_bytes).decode('utf-8')
                pdf_doc.close()
                print(f"[DEBUG] PDF转换完成: 第一页 -> PNG ({len(img_bytes)} 字节)")
            except Exception as pdf_e:
                print(f"[ERROR] PDF转换失败: {pdf_e}")
                return JSONResponse(
                    status_code=400,
                    content={"code": 400, "detail": "PDF文件处理失败，请尝试图片格式"}
                )
        else:
            # 普通图片
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            print("[DEBUG] ✅ Base64编码完成")
        
        # 调用百度 OCR
        print("[DEBUG] 正在调用百度OCR...")
        ocr_result = call_baidu_ocr(image_base64)
        
        if ocr_result is None:
            print("[ERROR] ❌ OCR返回空结果")
            return JSONResponse(
                status_code=500,
                content={"code": 500, "detail": "OCR服务无响应，请稍后重试"}
            )
        
        # 检查OCR错误
        if "error" in ocr_result:
            err_val = ocr_result.get("error", "未知错误")
            if isinstance(err_val, dict):
                error_msg = err_val.get("msg", str(err_val))
            else:
                error_msg = str(err_val)
            print(f"[ERROR] OCR返回错误: {error_msg}")
            return JSONResponse(
                status_code=500,
                content={"code": 500, "detail": f"OCR识别失败: {error_msg}"}
            )
        
        # 提取文字结果
        if "words_result" not in ocr_result:
            print(f"[WARN] OCR结果中没有words_result: {ocr_result}")
            return JSONResponse(
                status_code=500,
                content={"code": 500, "detail": "OCR识别失败，请确认图片包含清晰文字"}
            )
        
        words = [item["words"] for item in ocr_result.get("words_result", [])]
        full_text = "\n".join(words)
        print(f"[DEBUG] OCR识别到 {len(words)} 行文字")
        print(f"[DEBUG] 识别内容预览: {full_text[:200]}...")
        
        if not full_text.strip():
            print("[WARN] OCR识别结果为空")
            return JSONResponse(
                status_code=500,
                content={"code": 500, "detail": "未识别到任何文字，请确认发票图片清晰"}
            )
        
        # ===== 发票验证 =====
        # 检查是否包含发票关键词
        invoice_keywords = ["发票", "发票代码", "发票号码", "税务局", "税务"]
        is_invoice = any(kw in full_text for kw in invoice_keywords)
        
        print(f"[DEBUG] 发票验证: {'是发票' if is_invoice else '非发票'} | 关键词匹配: {[kw for kw in invoice_keywords if kw in full_text]}")
        
        # 非发票：返回特殊 code 让前端提示
        if not is_invoice and len(words) < 5:
            print("[WARN] OCR结果非发票")
            return JSONResponse(
                status_code=200,
                content={"code": 1, "detail": "上传的不是发票，请确认上传发票图片"}
            )
        
        # ===== 改进的发票信息提取 =====
        invoice_data = {
            "full_text": full_text,
            "number": "",
            "amount": "",
            "seller": "",
            "date": ""
        }
        
        # 合并所有文本为一行，方便搜索
        all_text = " ".join(words)
        all_text_clean = all_text.replace(" ", "").replace("\t", "")
        
        # 1. 发票号码：多种格式
        num_patterns = [
            r'发票号码?[：:]\s*(\w+)',
            r'发票代码[：:]\s*(\w+)',
            r'No[.：:]\s*(\w+)',
            r'号码[：:]\s*(\w+)',
            r'(\d{8}[\dXx]{4})',  # 12位发票号
        ]
        for pat in num_patterns:
            match = re.search(pat, all_text)
            if match:
                invoice_data["number"] = match.group(1)
                break
        
        # 2. 金额
        amount_patterns = [
            r'价税合计[\(（]大写[\)）]?[：:^\d]*[\(（]小写[\)）]?[：:]\s*[¥￥]?(\d+\.\d{2})',
            r'价税合计[：:]\s*[¥￥]?(\d+\.?\d*)',
            r'合计[：:]\s*[¥￥]?(\d+\.?\d*)',
            r'小写[\)）]?[：:]\s*[¥￥]?(\d+\.?\d*)',
            r'金额[：:]\s*[¥￥]?(\d+\.?\d*)',
            r'税额[：:]\s*[¥￥]?(\d+\.?\d*)',
            r'价税合计[\(（]小写[\)）]?[：:]*\s*[¥￥]?(\d+\.\d{2})',
            r'小写[\(（].*?[\)）][：:]*\s*[¥￥]?(\d+\.\d{2})',
        ]
        for pat in amount_patterns:
            match = re.search(pat, all_text)
            if match:
                invoice_data["amount"] = match.group(1)
                break
        
        # 3. 销售方（销方名称）
        seller_patterns = [
            r'(?:销售方|销方|收款人)[\(（]名称[\)）]?[：:]\s*(.+?)(?:\s{2,}|$)',
            r'(?:销售方|销方)[：:]\s*(.+?)(?:\s{2,}|$)',
            r'名称[：:]\s*(.+?公司)',
            r'(.+?有限公司)',
            r'(.+?有限责任公司)',
            r'(.+?集团)',
            r'(.+?厂)',
        ]
        for line in words:
            for pat in seller_patterns:
                match = re.search(pat, line)
                if match:
                    name = match.group(1).strip()
                    if len(name) > 2 and len(name) < 50:
                        invoice_data["seller"] = name
                        break
            if invoice_data["seller"]:
                break
        
        # 后备：任何包含"公司"的行
        if not invoice_data["seller"]:
            for line in words:
                if "公司" in line or "有限" in line:
                    invoice_data["seller"] = line.strip().replace(" ", "")
                    break
        
        # 4. 开票日期
        date_patterns = [
            r'开票日期[：:]\s*(\d{4}[-年.]\d{1,2}[-月.]\d{1,2})',
            r'日期[：:]\s*(\d{4}[-年.]\d{1,2}[-月.]\d{1,2})',
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]
        for pat in date_patterns:
            match = re.search(pat, all_text)
            if match:
                raw = match.group(0)
                invoice_data["date"] = raw.replace("年", "-").replace("月", "-").replace("日", "").replace(".", "-")
                break
        
        # 5. 如果以上字段提取不完整，尝试用 DeepSeek 补全
        missing_fields = [k for k in ["number", "amount", "seller", "date"] if not invoice_data[k]]
        if missing_fields:
            print(f"[DEBUG] 字段缺失({len(missing_fields)}个): {missing_fields}，尝试DeepSeek补全...")
            try:
                ds_prompt = f"""从以下发票OCR识别文本中提取关键信息，只返回JSON格式，不要多余文字：

ocr文本：
{full_text}

请提取：
1. number: 发票号码
2. amount: 金额（纯数字）
3. seller: 销售方名称
4. date: 开票日期

返回格式：{{"number":"","amount":"","seller":"","date":""}}
"""
                ds_result = call_deepseek(ds_prompt)
                if ds_result:
                    # 尝试从DeepSeek结果中提取JSON
                    json_match = re.search(r'\{[^{}]*\}', ds_result)
                    if json_match:
                        ds_data = json.loads(json_match.group())
                        for k in missing_fields:
                            if ds_data.get(k) and not invoice_data[k]:
                                invoice_data[k] = ds_data[k]
                                print(f"[DEBUG] DeepSeek补全 {k}={ds_data[k]}")
            except Exception as ds_err:
                print(f"[DEBUG] DeepSeek补全省略: {ds_err}")
        
        # 提取品名（第一行有意义的内容）
        purpose = ""
        for line in words:
            line_s = line.strip()
            if len(line_s) >= 4 and not any(kw in line_s for kw in ["发票", "代码", "号码", "日期", "密码", "税控", "校验", "收款", "复核", "开票", "¥", "合计", "小写", "大写", "销售方", "购买方", "纳税人"]):
                purpose = line_s[:50]
                break
        
        print(f"[DEBUG] ✅ 提取结果: 号码={invoice_data['number']}, 金额={invoice_data['amount']}, 销方={invoice_data['seller'][:20] if invoice_data.get('seller') else ''}, 日期={invoice_data['date']}, 品名={purpose[:20]}")
        
        return {
            "code": 0,
            "data": {
                **invoice_data,
                "purpose": purpose
            }
        }
        
    except Exception as e:
        print(f"[ERROR] OCR接口异常: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={"code": 500, "detail": f"服务器内部错误: {str(e)}"}
        )

# ================= 金额转大写 =================
def amount_to_chinese(n):
    """数字金额转中文大写，如 1280.50 → 壹仟贰佰捌拾元伍角"""
    units = ["", "拾", "佰", "仟", "万", "拾", "佰", "仟", "亿"]
    digits = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
    try:
        n = round(float(n), 2)
    except:
        return "零元整"
    integer_part = int(n)
    decimal_part = round((n - integer_part) * 100)
    
    if integer_part == 0 and decimal_part == 0:
        return "零元整"
    
    # 处理整数部分
    int_str = str(integer_part)
    result = ""
    zero_flag = False
    need_unit = True
    
    for i, ch in enumerate(reversed(int_str)):
        if ch != '0':
            if zero_flag:
                result = "零" + result
                zero_flag = False
            unit = units[i] if i < len(units) else ""
            result = digits[int(ch)] + unit + result
        else:
            zero_flag = True
    
    result += "元"
    
    # 小数部分
    if decimal_part == 0:
        result += "整"
    else:
        jiao = decimal_part // 10
        fen = decimal_part % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        if fen > 0:
            result += digits[fen] + "分"
    
    return result

@app.post("/api/generate_word")
async def generate_word(
    department: str = Form(...),
    operator: str = Form(...),
    category: str = Form(""),
    amount: str = Form(...),
    remark: str = Form(""),
    report_date: str = Form(...),
    invoice_number: str = Form(""),
    seller: str = Form(""),
    leader_approval: str = Form(""),
    dept_audit: str = Form(""),
    accountant: str = Form(""),
    cashier: str = Form(""),
    original_loan: str = Form(""),
    refund: str = Form(""),
    items_json: str = Form("[]")
):
    """根据模板生成报销单 Word 文档"""
    try:
        import json
        template_path = os.path.join(os.path.dirname(__file__), "模板_费用报销单.docx")
        
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            return JSONResponse(status_code=500, content={"code": 500, "detail": "模板文件缺失"})
        
        items = json.loads(items_json) if items_json else []
        
        # ===== 填充段落 =====
        for p in doc.paragraphs:
            text = p.text
            
            # P1: 报销部门 + 日期 + 附件页
            if "报销部门" in text and "单据及附件" in text:
                for run in p.runs:
                    run.text = ""
                date_str = report_date.replace("-", " 年 ") + " 月 " + " 日 "
                p.runs[0].text = f"报销部门：{department}   {date_str}填                单据及附件    {len(items)}    页"
                break
        
        for p in doc.paragraphs:
            text = p.text
            # P2: 会计主管 会计 出纳 报销人 领款人
            if "会计主管" in text and "领款人" in text:
                for run in p.runs:
                    run.text = ""
                p.runs[0].text = f"会计主管  {accountant}        会计  {cashier}        出纳          报销人  {operator}            领款人  {operator}"
                break
        
        # ===== 填充表格 =====
        if doc.tables:
            tbl = doc.tables[0]
            
            # Rows 1-2 (数据行)
            for i in range(min(len(items), 2)):
                row_idx = i + 1  # Row 1 or 2
                if row_idx < len(tbl.rows):
                    row = tbl.rows[row_idx]
                    row.cells[0].text = items[i].get("purpose", "")
                    row.cells[1].text = str(items[i].get("amount", ""))
                    row.cells[2].text = items[i].get("remark", "")
            
            # Row 5 (合计行)
            if len(tbl.rows) > 5:
                tbl.rows[5].cells[0].text = "合    计"
                tbl.rows[5].cells[1].text = str(amount)
            
            # Row 6 (金额大写/原借款/应退余款)
            if len(tbl.rows) > 6:
                tbl.rows[6].cells[0].text = "金额大写：" + amount_to_chinese(amount)
                tbl.rows[6].cells[1].text = "金额大写：" + amount_to_chinese(amount)
                tbl.rows[6].cells[2].text = f"原借款：{original_loan}" if original_loan else "原借款："
                tbl.rows[6].cells[3].text = f"原借款：{original_loan}" if original_loan else "原借款："
                tbl.rows[6].cells[4].text = f"应退余款：{refund}" if refund else "应退余款："
                tbl.rows[6].cells[5].text = f"应退余款：{refund}" if refund else "应退余款："
            
            # 发票凭证表 (Table 1)
            if len(doc.tables) > 1:
                inv_tbl = doc.tables[1]
                lines = [f"{i.get('purpose','?')} ¥{i.get('amount','0')}" for i in items]
                inv_tbl.cell(0, 0).text = "发票凭证：" + "；".join(lines) if lines else "发票凭证"
        
        filename = f"费用报销单_{operator}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join("generated", filename)
        os.makedirs("generated", exist_ok=True)
        doc.save(filepath)
        
        return {"code": 0, "data": {"filename": filename, "url": f"/download/{filename}"}}
    except Exception as e:
        print(f"[ERROR] 生成Word错误: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{filename}")
async def download_file(filename: str):
    filepath = os.path.join("generated", filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(filepath, filename=filename)

# ================= 调试端点 =================
@app.post("/api/debug_ocr")
async def debug_ocr(file: UploadFile = File(...)):
    """调试OCR，返回原始识别内容"""
    import json
    image_data = await file.read()
    image_base64 = base64.b64encode(image_data).decode('utf-8')
    token = get_baidu_ocr_token()
    if not token:
        return {"error": "无法获取百度Token"}
    
    url = f"https://aip.baidubce.com/rest/2.0/ocr/v1/general_basic?access_token={token}"
    resp = requests.post(url, data={"image": image_base64}, headers={"content-type": "application/x-www-form-urlencoded"}, timeout=30)
    
    result = resp.json()
    words = [item["words"] for item in result.get("words_result", [])]
    
    return {
        "baidu_raw": result,
        "words": words,
        "full_text": "\n".join(words)
    }

# ================= 启动服务 =================
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print("=" * 50)
    print("🚀 启动数智办公后端服务...")
    print(f"📡 访问地址: http://0.0.0.0:{port}")
    print(f"📖 API文档: http://0.0.0.0:{port}/docs")
    print(f"📂 generated目录: {os.path.abspath('generated')}")
    print("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=port)